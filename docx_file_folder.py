import docx
import os

class DocxFileFolder:

    def __init__(self,file_path):
        self.file_path = file_path
        global doc
        doc = docx.Document(file_path)


    def validate_docx_file(self):
        first_validation = ["Job Names:", "Folder Name:","Description:"]
        forth_table_validation = ["Alerts:","submitted by","greater than:","not finished by"]
        try:
            if doc.tables:
                if all(x in doc.tables[0].rows[1].cells[0].text for x in first_validation):
                    if all(x in doc.tables[3].rows[1].cells[0].text for x in forth_table_validation):
                        return True
            return False
        except IndexError:
            return False


    def __get_general_info_list(self):
        info_list = []
        for paragraphs in doc.tables[0].rows[1].cells[0].paragraphs:
            if paragraphs.text == "" or paragraphs.text.isspace():
                continue
            info_list.append(paragraphs.text)
        return info_list


    def get_folder_name(self):
        return self.__get_general_info_list()[0][12:].strip()


    def get_description(self):
        return self.__get_general_info_list()[1][12:].strip()


    def get_job_names(self):
        job_names = []
        for rows in doc.tables[0].rows:
            if rows.cells[0].text.isdigit():
                job_names.append(rows.cells[1].text)
        return job_names


    def __get_schedule_info_list(self):
        info_list = []
        for paragraphs in doc.tables[1].rows[1].cells[0].paragraphs:
            if paragraphs.text == "" or paragraphs.text.isspace():
                continue
            info_list.append(paragraphs.text)
        return info_list


    def get_calendar(self):
        return self.__get_schedule_info_list()[0][9:].strip()


    def get_from_time(self):
        return self.__get_schedule_info_list()[1][10:].strip()


    def get_to_time(self):
        return self.__get_schedule_info_list()[2][8:].strip()


    def get_run_frequency(self):
        return self.__get_schedule_info_list()[3][14:].strip()


    def get_in_condition(self):
        in_condition = []
        for rows in doc.tables[2].rows:
            if rows.cells[0].text.isdigit():
                in_condition.append(rows.cells[1].text)
        return in_condition


    def __get_action_list(self):
        action_list = []
        for paragraphs in doc.tables[3].rows[1].cells[0].paragraphs:
            if paragraphs.text == "" or paragraphs.text == "Alerts:" or paragraphs.text.isspace():
                continue
            action_list.append(paragraphs.text)
        return action_list


    def get_not_submitted(self):
        return self.__get_action_list()[0][29:].strip()


    def get_execution_time(self):
        return self.__get_action_list()[1][45:].strip()


    def get_not_finished(self):
        return self.__get_action_list()[2][28:].strip()


    def file_basename(self):
        return os.path.basename(self.file_path)