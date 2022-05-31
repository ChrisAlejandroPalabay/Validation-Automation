import docx
import os

class DocxFileJob:

    def __init__(self,file_path):
        self.file_path = file_path
        global doc
        doc = docx.Document(file_path)


    def validate_docx_file(self):
        first_validation = ["Job Name:", "Folder Name:","Description:","Command (Script) Name:"]
        forth_table_validation = ["Alerts:","submitted by","greater than:","not finished by"]
        try:
            if doc.tables:
                if all(x in doc.tables[0].rows[1].cells[0].text for x in first_validation):
                    if all(x in doc.tables[3].rows[1].cells[0].text for x in forth_table_validation):
                        return True
            return False
        except IndexError:
            return False


    def __get_general_info_list(self):
        info_list = []
        for paragraphs in doc.tables[0].rows[1].cells[0].paragraphs:
            if paragraphs.text == "" or paragraphs.text.isspace():
                continue
            info_list.append(paragraphs.text)
        return info_list


    def get_job_name(self):
        return self.__get_general_info_list()[0][9:].strip()


    def get_folder_name(self):
        return self.__get_general_info_list()[1][12:].strip()


    def get_description(self):
        return self.__get_general_info_list()[2][12:].strip()


    def get_command_prd(self):
        return self.__get_general_info_list()[4][9:].strip()


    def get_command_nonprd(self):
        return self.__get_general_info_list()[5][13:].strip()


    def __get_parameters_info(self):
        params = []
        for row in doc.tables[0].rows:
            if row.cells[0].text.isdigit():
                params.append(row.cells[1].text)
            continue
        return params

    def get_param_prd(self):
        return self.__get_parameters_info()[0]

    def get_param_nonprd(self):
        return self.__get_parameters_info()[1]


    def __prd_nonprd_info(self):
        prd_info = []
        for paragraphs in doc.tables[0].rows[6].cells[0].paragraphs:
            if paragraphs.text == "" or paragraphs.text.isspace():
                continue
            prd_info.append(paragraphs.text)
        return prd_info


    def __get_prd_info(self):
        prd_info = self.__prd_nonprd_info()[0][33:].strip()
        return prd_info.split("|")


    def get_prd_hostgroup(self):
        return self.__get_prd_info()[0]


    def get_prd_hostname(self):
        return self.__get_prd_info()[1]


    def get_prd_runas(self):
        return self.__get_prd_info()[2]


    def __get_nonprd_info(self):
        nonprd_info = self.__prd_nonprd_info()[1][37:].strip()
        return nonprd_info.split("|")


    def get_nonprd_hostgroup(self):
        return self.__get_nonprd_info()[0]


    def get_nonprd_hostname(self):
        return self.__get_nonprd_info()[1]


    def get_nonprd_runas(self):
        return self.__get_nonprd_info()[2]


    def __get_schedule_info_list(self):
        schedule_info = []
        for paragraphs in doc.tables[1].rows[1].cells[0].paragraphs:
            if paragraphs.text == "" or paragraphs.text.isspace():
                continue
            schedule_info.append(paragraphs.text)
        return schedule_info


    def get_calendar(self):
        return self.__get_schedule_info_list()[0][9:].strip()


    def get_from_time(self):
        return self.__get_schedule_info_list()[1][10:].strip()


    def get_to_time(self):
        return self.__get_schedule_info_list()[2][8:].strip()


    def get_cyclic(self):
        return self.__get_schedule_info_list()[3][7:].strip()


    def get_run_frequency(self):
        return self.__get_schedule_info_list()[4][14:].strip()


    def get_in_condition(self):
        in_condition = []
        for rows in doc.tables[2].rows:
            if rows.cells[0].text.isdigit():
                in_condition.append(rows.cells[1].text)
        return in_condition


    def __get_action_list(self):
        action_list = []
        for paragraphs in doc.tables[3].rows[1].cells[0].paragraphs:
            if paragraphs.text == "" or paragraphs.text == "Alerts:" or paragraphs.text.isspace():
                continue
            action_list.append(paragraphs.text)
        return action_list


    def get_not_submitted(self):
        return self.__get_action_list()[0][26:].strip()


    def get_execution_time(self):
        return self.__get_action_list()[1][43:].strip()


    def get_not_finished(self):
        return self.__get_action_list()[2][25:].strip()

    # For testing purposes
    def test_run(self):
        for items in self.get_in_condition():
            print(items)

    def file_basename(self):
        return os.path.basename(self.file_path)
    